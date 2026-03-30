from fpdf import FPDF
import io
import pandas as pd
from docx import Document

def export_to_format(df: pd.DataFrame, file_format: str, title: str = "Analysis Report"):
    """
    Unified export function supporting CSV, Excel, DOCX, and PDF.
    Returns bytes in requested format.
    """
    if file_format == "CSV":
         return df.to_csv(index=False).encode('utf-8')
    elif file_format == "Excel":
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Report')
        return output.getvalue()
    elif file_format == "DOCX":
        doc = Document()
        doc.add_heading(title, 0)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns): hdr_cells[i].text = str(col)
        # Limit rows for docx performance in demo
        for _, row in df.head(100).iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row): row_cells[i].text = str(val)
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    elif file_format == "PDF":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("helvetica", "B", 16)
        pdf.cell(0, 10, title, ln=True, align="C")
        pdf.ln(5)
        pdf.set_font("helvetica", "B", 8)
        cols = list(df.columns)[:8] # Cap columns to prevent spill
        w = pdf.epw / len(cols)
        for c in cols: pdf.cell(w, 8, str(c)[:15], border=1) # Truncate header
        pdf.ln()
        pdf.set_font("helvetica", "", 7)
        for _, r in df.head(100).iterrows(): # Limit rows for browser responsiveness
             for c in cols: pdf.cell(w, 7, str(r[c])[:25], border=1)
             pdf.ln()
        return bytes(pdf.output())
    return None
